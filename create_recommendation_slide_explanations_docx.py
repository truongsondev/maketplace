from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "Giai_Thich_Chi_Tiet_Tung_Slide_AI_Recommendation_Aura.docx"


slides = [
    ("AI Recommendation", "Mở đầu bằng mục tiêu của bài: không mô tả một endpoint trả danh sách sản phẩm, mà giải thích toàn bộ chuỗi quyết định khiến một sản phẩm đứng hạng nhất. Hệ thống của Aura Fashion kết hợp embedding 384 chiều, điểm hybrid, bốn loại feed và nhiều nguồn dữ liệu. Cần nhấn mạnh đây là kiến trúc đang triển khai trong dự án: Node.js làm orchestration, FastAPI đảm nhiệm semantic retrieval, PostgreSQL/pgvector lưu vector, còn MySQL, Redis và RabbitMQ phục vụ dữ liệu hành vi và vận hành."),
    ("Câu hỏi trung tâm: vì sao sản phẩm này được gợi ý?", "Giải thích ba lớp bằng chứng. Thứ nhất là hành vi: xem, yêu thích, thêm giỏ, mua và truy vấn tìm kiếm. Thứ hai là nội dung sản phẩm: tên, mô tả, danh mục và thuộc tính được biến thành vector. Thứ ba là quyết định xếp hạng: candidate từ nhiều nguồn được hợp nhất, lọc và sắp xếp. Vì vậy câu trả lời đúng không phải “do API trả về”, mà là sản phẩm có tổng bằng chứng mạnh nhất tại placement và thời điểm hiện tại."),
    ("Kiến trúc triển khai thực tế", "Đi từ trái sang phải theo ba luồng. Online path phục vụ request: kiểm tra Redis, nếu miss thì lấy candidate từ SQL/Redis và AI, merge, tạo product card rồi cache. Near-real-time path xử lý event qua RabbitMQ để ghi MySQL, tăng hot score và xóa cache liên quan. Offline path chạy định kỳ để tính co-occurrence, popularity, preference và embedding. Node.js là tầng điều phối và luật nghiệp vụ; AI service chỉ là một nguồn retrieval/scoring, nên AI lỗi không nhất thiết làm cả feed lỗi."),
    ("Dữ liệu đầu vào: một event không chỉ là một click", "Giải thích ý nghĩa từng trường của RecommendationEvent. eventType cho biết mức độ quan tâm; userId phục vụ cá nhân hóa người đăng nhập; sessionId giữ ngữ cảnh khách ẩn danh; productId là đối tượng; searchQuery thể hiện ý định; placement giúp đo hiệu quả từng vị trí; occurredAt phục vụ recency; dedupeKey bảo đảm một sự kiện không bị ghi nhiều lần. Một event tốt vừa dùng để học sở thích, vừa hỗ trợ đo lường và vận hành tin cậy."),
    ("Ingestion đáng tin cậy: event đi qua RabbitMQ", "UI chỉ publish message nên không phải chờ toàn bộ quá trình ghi dữ liệu và cập nhật feature. Consumer đọc với prefetch 20, persist theo dedupeKey, cập nhật Redis rồi invalidates cache. Nếu xử lý lỗi, message được chuyển sang retry queue có TTL 15 giây; tối đa ba lần rồi vào dead-letter queue. Message durable và delivery mode 2 giúp giảm nguy cơ mất dữ liệu. Idempotency đặc biệt quan trọng vì retry không được phép nhân đôi hot score hoặc lịch sử hành vi."),
    ("Trọng số hành vi", "Các hành vi không có giá trị như nhau: purchase là bằng chứng mạnh nhất, tiếp theo là add-to-cart, favorite và view. Personalized còn boost 35% tín hiệu thuộc chính user. Cần nói rõ các con số hiện tại là heuristic mã hóa hiểu biết kinh doanh, không phải trọng số được học. Chúng dễ giải thích và khởi động nhanh, nhưng phải được kiểm nghiệm bằng dữ liệu thực tế và A/B testing để tránh thiên lệch chủ quan."),
    ("Offline refresh", "Scheduler chạy khi hệ thống khởi động và lặp khoảng 15 phút. Nó đọc tối đa 2.000 sản phẩm, 5.000 event gần đây, tính cặp sản phẩm cùng đơn bằng self-join, gọi AI train để tạo preference/popularity và gọi embed/products để upsert vector. Kết quả là các artifact đã tính trước, giúp request online không phải xử lý toàn bộ lịch sử. Đánh đổi là dữ liệu semantic/co-occurrence có độ trễ theo chu kỳ refresh."),
    ("Tạo document cho một sản phẩm", "Minh họa bằng váy ren Rumy: tên, mô tả, danh mục, màu, dịp sử dụng và chất liệu được ghép thành một chuỗi. Chuỗi đi qua normalize, tokenizer, encoder, L2 normalization rồi lưu vector và metadata. Embedding chỉ hiểu những gì được đưa vào document; nếu thiếu màu, chất liệu hoặc variant attributes, truy vấn semantic sẽ thiếu bằng chứng. Đây là lý do chất lượng pipeline dữ liệu quan trọng không kém lựa chọn model."),
    ("Encoder: MiniLM nếu tải được, hashing nếu không", "all-MiniLM-L6-v2 tạo vector 384 chiều và có khả năng đặt các câu đồng nghĩa gần nhau dù không trùng token. Khi model không tải được, hashing-baseline băm từng token vào 384 ô rồi chuẩn hóa. Hashing nhanh, chạy offline nhưng không hiểu ngữ cảnh và có collision. Fallback này tăng availability, tuy nhiên health endpoint phải báo model_loaded và monitoring cần cảnh báo, vì dịch vụ vẫn trả 200 trong khi chất lượng đã suy giảm."),
    ("Cosine similarity trong pgvector", "Cosine đo góc giữa vector query và vector sản phẩm. Vì các vector đã chuẩn hóa L2, cosine gần bằng tích vô hướng; score SQL được tính bằng 1 trừ khoảng cách cosine của toán tử <=>. Điểm gần 1 biểu thị nội dung rất gần, điểm thấp biểu thị ít liên quan. IVFFlat chia không gian vector để tìm nhanh trong một phần dữ liệu thay vì full scan, nhưng cần ANALYZE và tinh chỉnh lists/probes để cân bằng tốc độ với recall."),
    ("Context vector", "Hệ thống không tạo user embedding cố định trong mỗi request. Nó lấy các sản phẩm người dùng tương tác gần đây, đọc vector của chúng, tính trung bình rồi chuẩn hóa để biểu diễn “gu hiện tại”. Ví dụ trung bình váy ren, váy chữ A và áo cổ bèo tạo context thiên về phong cách nữ tính. Mean vector đơn giản và nhanh nhưng coi các item gần như ngang nhau; tương lai có thể dùng trọng số hành vi và recency. Nếu không có vector context, semantic hits rỗng và popularity trở thành fallback chính."),
    ("Công thức hybrid trong AI service", "Trình bày công thức final = 0,75 × cosine + 0,25 × log(1 + popularity). Cosine đại diện relevance nội dung, popularity đại diện demand thực tế; log làm giảm mức tăng của sản phẩm quá nổi tiếng. Tuy nhiên log1p(popularity) vẫn có thể lớn hơn 1 trong khi cosine nằm khoảng 0–1. Bảng ví dụ cho thấy váy basic có semantic thấp hơn nhưng thắng nhờ popularity. Do đó tỷ lệ 75/25 không thực sự có ý nghĩa nếu chưa chuẩn hóa hai thành phần về cùng thang."),
    ("Candidate generation", "Không có một model duy nhất quyết định tất cả. pgvector tìm sản phẩm gần về ngữ nghĩa; co-occurrence học sản phẩm mua cùng; hot và trending phản ánh nhu cầu; search intent phản ánh câu tìm kiếm; category/type và latest/sale cứu cold-start. Mỗi nguồn trả productId, score, reason và source. Hệ thống lấy nhiều hơn limit yêu cầu để ưu tiên recall, sau đó mới hợp nhất, loại trùng, áp guardrail và cắt top-K."),
    ("Feed Home", "Home phải hữu ích cả với khách chưa đăng nhập. Nó hợp nhất hot score thời gian thực, trending 30 ngày, truy vấn của session và hành vi session 14 ngày. Add-to-cart có trọng số cao hơn view vì thể hiện ý định mua mạnh hơn. Cache key gắn session và limit, TTL 900 giây. Khi không có context, fallback chọn sản phẩm có ảnh, ưu tiên đang sale và mới tạo. Đây là sự kết hợp giữa khám phá xu hướng và cá nhân hóa tức thời."),
    ("Feed Product Detail", "Context chỉ là sản phẩm hiện tại. AI tìm vector gần, co-occurrence tìm item thường được mua cùng và category/type cung cấp fallback. Sau merge, sản phẩm hiện tại bắt buộc bị exclude để tránh gợi ý chính nó. Semantic hữu ích cho sản phẩm mới chưa có giao dịch; co-occurrence mạnh khi đã có đơn hàng thật; category bảo đảm feed không rỗng. Ba nguồn bổ sung cho nhau thay vì cạnh tranh như các model riêng biệt."),
    ("Feed Cart", "Toàn bộ sản phẩm trong giỏ vừa là context vừa là exclusion. Vector trung bình có thể mang nhiều ý định, chẳng hạn váy ren cộng túi trắng tạo tín hiệu về phong cách và dịp sử dụng; giày cao gót có thể gần cả hai. Co-occurrence cung cấp cross-sell từ đơn hàng thật, category mở rộng họ sản phẩm. ADD/REMOVE phải invalidates cache để request tiếp theo phản ánh giỏ mới. Feed không được trả lại bất kỳ item nào đã có trong giỏ."),
    ("Feed Personalized", "Pipeline lấy hành vi 45 ngày, tính điểm, tạo context và đồng thời loại context, item trong giỏ và sản phẩm đã mua 90 ngày. Candidate đến từ AI, search intent, related pairs, category và catalog fallback. Điểm đáng chú ý là own signals được đưa lại làm candidate nhưng contextIds lại nằm trong exclusion, nên nhiều candidate từ chính lịch sử gần đây bị loại. Quy tắc chống lặp tốt cho khám phá nhưng có thể không phù hợp với sản phẩm cần mua lại định kỳ."),
    ("Search intent", "Hệ thống gom query trong 30 ngày, xếp theo số lần và độ mới, rồi dùng SQL LIKE trên name/description. Công thức giảm điểm theo vị trí query và vị trí sản phẩm, tối thiểu 0,5. Ưu điểm là nhanh, minh bạch và hoạt động tốt với cụm rõ như “váy trắng”. Nhược điểm là không hiểu “đầm dự tiệc” gần “váy đi tiệc”, không xử lý typo tốt và score count×3 không cùng thang với cosine. Query embedding là hướng nâng cấp hợp lý."),
    ("Co-occurrence", "Self-join order_items trên cùng orderId đếm số lần hai productId xuất hiện chung. Nếu A và B cùng 12 đơn thì score 12, cao hơn A và C cùng 4 đơn. Đây là collaborative signal đơn giản, dễ kiểm tra và không cần mô hình phức tạp. Điểm yếu là bestseller xuất hiện trong nhiều đơn nên có thể liên quan giả với mọi sản phẩm. Có thể cải thiện bằng cosine/Jaccard, lift, PMI hoặc chuẩn hóa theo tần suất biên."),
    ("Merge scores trong Node", "Khi cùng product xuất hiện từ nhiều nguồn, thuật toán giữ score lớn nhất; mỗi nguồn phụ đến sau đóng góp 20% score. Cách này thưởng cho candidate có nhiều bằng chứng. Nhưng cosine khoảng 0–1, co-occurrence là số đếm, search dựa count×3 và category dựa thứ hạng limit…1. Vì khác thang, một category score 8 dễ lấn át AI score 1,085. Trước khi tối ưu trọng số, phải calibrate hoặc chuyển sang rank fusion như RRF."),
    ("Bộ lọc và guardrail", "Sau ranking, hệ thống loại item đang xem, trong giỏ, đã mua gần đây, soft-deleted, thiếu product card và loại bản sao productId. Những luật này bảo đảm kết quả hợp lệ về ngữ cảnh. Khoảng trống hiện tại là card lookup chưa bảo đảm Product.status ACTIVE, variant đang hoạt động và stockAvailable lớn hơn 0. Nếu không bổ sung, một candidate điểm cao nhưng hết hàng vẫn có thể xuất hiện, làm giảm trải nghiệm và conversion."),
    ("Worked example #1", "Tính từng cột cho ba candidate. B có cosine 0,91 nhưng popularity 4 nên AI score 1,085. C có cosine 0,72 nhưng popularity 20, log1p bằng 3,045 nên đạt 1,301 và đứng đầu. D có semantic khá nhưng popularity thấp nên cuối. Ví dụ chứng minh “semantic weight 75%” chưa đồng nghĩa semantic chắc chắn chi phối, vì đầu vào popularity chưa chuẩn hóa. Có thể dùng percentile, min-max hoặc capped/log-normalized score."),
    ("Worked example #2", "Tiếp tục ví dụ qua tầng Node. C nhận category score 8 và AI 1,301 nên merge thành 8,260; D nhận category 7, co-occurrence 1,2 và AI 0,796 nên đạt khoảng 7,399; B chỉ có AI và co-occurrence nên xuống hạng ba. Cần phân biệt semantic winner, AI hybrid winner và displayed winner. Kết quả cuối là sản phẩm thắng sau orchestration, không đơn thuần là output của vector model."),
    ("End-to-end", "Theo dấu một candidate từ catalog đến #1: metadata tạo document; MiniLM sinh vector; các vector hành vi tạo context; pgvector lấy top candidate; AI kết hợp cosine và popularity; Node merge với category/co-occurrence; guardrail quyết định nó có được hiển thị không. Ví dụ C đi từ cosine 0,72, nhờ popularity thành 1,301, nhờ category thành 8,260 và không bị exclude nên đứng đầu. Chuỗi này giúp giải thích recommendation bằng dữ liệu cụ thể."),
    ("Cache hai tầng", "Request đầu tiên kiểm tra Redis; nếu miss thì đọc RecommendationCache trong MySQL; DB hit sẽ warm lại Redis; full miss mới chạy toàn pipeline và persist kết quả. Redis cho tốc độ, DB cache cho khả năng phục hồi khi Redis mất dữ liệu. TTL thay đổi theo feed: product detail lâu hơn vì context ít đổi, cart ngắn hơn vì hành vi thay đổi nhanh. Cache key phải chứa đúng user/session/product/limit để tránh trả nhầm context."),
    ("Cache invalidation", "Freshness không chỉ dựa vào TTL. Event có sessionId xóa home và personalized của session; event có userId xóa cart và personalized của user. Vì vậy một lần add-to-cart có thể làm feed đổi ngay ở request kế tiếp. Vấn đề kỹ thuật là Redis KEYS(pattern) quét toàn keyspace và có thể block production. Nên dùng SCAN, lưu tập key liên quan hoặc dùng version number trong key để invalidation theo O(1)."),
    ("Fault tolerance", "AI request có thể timeout hoặc trả 5xx. Repository bắt exception, tăng metric fallback, coi AI candidates là mảng rỗng rồi tiếp tục với co-occurrence, category và catalog. RabbitMQ retry/DLQ bảo vệ ingestion; DB cache hỗ trợ khi Redis không có dữ liệu. Đây là graceful degradation: chất lượng có thể giảm nhưng endpoint vẫn phục vụ. Blind spot là fetch sang AI cần explicit timeout/AbortController, circuit breaker và giới hạn retry để tránh request bị treo."),
    ("Observability", "Các metric hiện có trả lời: event có đi vào không, cache hit bao nhiêu, feed mất bao lâu, AI operation chậm thế nào và fallback xảy ra mấy lần. Chỉ biết endpoint 200 là chưa đủ; cần quan sát candidate count theo source, zero-result rate, score distribution và model version. Sau đó nối impression, click, cart và purchase để biết hệ thống không chỉ chạy nhanh mà còn tạo giá trị kinh doanh."),
    ("Đánh giá chất lượng", "Offline metric dùng lịch sử làm ground truth: Recall@K đo khả năng lấy đúng item, NDCG@K tính cả vị trí, MRR nhìn item đúng đầu tiên và coverage đo độ phủ catalog. Online metric đo CTR, add-to-cart, assisted purchase, conversion và doanh thu/session. Guardrail theo dõi P95/P99, tỷ lệ gợi ý hết hàng, diversity, novelty và popularity concentration. Một model tốt phải cân bằng relevance, trải nghiệm, tốc độ và mục tiêu kinh doanh."),
    ("A/B testing đúng cách", "Dùng stable hash của userId hoặc sessionId để người dùng luôn vào cùng bucket. Control chạy heuristic hiện tại; treatment có thể chuẩn hóa score, RRF, diversity hoặc stock guardrail. Mỗi impression phải ghi experiment và variant, sau đó liên kết click/cart/purchase. So sánh uplift cần đủ cỡ mẫu và kiểm định thống kê, đồng thời theo dõi guardrail. Schema experiment đã tồn tại nhưng assignment và logging trong serving cần hoàn thiện."),
    ("Technical debt", "Trình bày trung thực các giới hạn: score chưa calibrate; body profile được gửi nhưng chưa dùng; user embedding có schema nhưng serving chưa đọc; thiếu lọc status/stock; Python hash thay đổi giữa process; IVFFlat chưa tune; own signal xung đột exclusion; query search chỉ LIKE. Điểm tích cực là các lớp đã tách rõ, nên có thể thay từng heuristic hoặc model mà không phải viết lại toàn bộ kiến trúc."),
    ("Roadmap kỹ thuật", "Thứ tự ưu tiên dựa trên impact và rủi ro. P0 bổ sung ACTIVE và stock > 0. P1 chuẩn hóa điểm từng nguồn. P2 dùng Reciprocal Rank Fusion để giảm phụ thuộc thang score. P3 thêm MMR hoặc category cap để tăng diversity. P4 khi tracking đủ sạch mới huấn luyện Learning-to-Rank. Trong ngắn hạn ưu tiên an toàn production và đo lường; sau đó A/B query embedding, diversity và cuối cùng LightGBM/XGBoost."),
    ("Kịch bản demo kỹ thuật", "Demo không nên chỉ mở carousel. Đầu tiên phát VIEW/ADD và kiểm tra bảng event cùng Redis hot score. Tiếp theo chạy refresh và xem co-occurrence/vector. Truy vấn pgvector để so cosine ba candidate. Gọi feed và đọc score, reason, source. Gọi lần hai để chứng minh cache hit. Cuối cùng phát ADD_TO_CART mới, xác nhận cache bị invalidates và ranking đổi. Mỗi bước cung cấp bằng chứng cho một tầng kiến trúc."),
    ("Kết luận", "Chốt lại ba lớp. Retrieval dùng embedding, pgvector, co-occurrence và fallback để tạo tập ứng viên. Ranking dùng hybrid 0,75/0,25 rồi merge max+20%, exclusion và guardrail để chọn top-K. Learning loop đưa event qua RabbitMQ vào MySQL/Redis, refresh artifact và invalidates cache để recommendation thay đổi. Thông điệp cuối: sản phẩm đứng số 1 vì có bằng chứng tốt nhất sau toàn pipeline, và mọi bước đều có thể đo, kiểm tra, giải thích và cải tiến."),
]


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.3)
section.right_margin = Cm(2.3)

doc.styles["Normal"].font.name = "Aptos"
doc.styles["Normal"].font.size = Pt(11)
doc.styles["Normal"].paragraph_format.space_after = Pt(7)
doc.styles["Title"].font.name = "Aptos Display"
doc.styles["Title"].font.color.rgb = RGBColor(29, 53, 87)
doc.styles["Heading 1"].font.color.rgb = RGBColor(67, 97, 238)
doc.styles["Heading 2"].font.color.rgb = RGBColor(29, 53, 87)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AURA FASHION")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(67, 97, 238)

p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("GIẢI THÍCH CHI TIẾT TỪNG SLIDE\nAI RECOMMENDATION")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Tài liệu lời thuyết trình cho 34 slide")
r.italic = True
r.font.size = Pt(13)

doc.add_page_break()
doc.add_heading("Hướng dẫn sử dụng", level=1)
doc.add_paragraph(
    "Mỗi mục dưới đây tương ứng với một slide. Phần “Lời giải thích” có thể dùng trực tiếp làm lời nói khi thuyết trình. "
    "Phần “Điểm cần nhấn mạnh” giúp chốt ý trước khi chuyển slide."
)

for index, (title, explanation) in enumerate(slides, 1):
    doc.add_heading(f"Slide {index:02d} — {title}", level=1)
    label = doc.add_paragraph()
    run = label.add_run("Lời giải thích")
    run.bold = True
    run.font.color.rgb = RGBColor(67, 97, 238)
    doc.add_paragraph(explanation)

    emphasis = doc.add_paragraph()
    run = emphasis.add_run("Điểm cần nhấn mạnh: ")
    run.bold = True
    if index == 5:
        emphasis.add_run("RabbitMQ giúp xử lý bất đồng bộ, retry an toàn và không để việc ghi event làm chậm trải nghiệm người dùng.")
    elif index in (12, 20, 22, 23):
        emphasis.add_run("Phải so sánh các score trên cùng thang đo; nếu không, tỷ trọng và thứ hạng có thể khác với ý nghĩa thiết kế.")
    elif index in (21, 31, 32):
        emphasis.add_run("Chất lượng recommendation không chỉ là relevance; tính hợp lệ, tồn kho, khả năng đo lường và độ an toàn production cũng là yêu cầu cốt lõi.")
    elif index == 34:
        emphasis.add_run("Recommendation là một chuỗi quyết định có thể kiểm chứng, không phải một hộp đen hay một lời gọi API đơn lẻ.")
    else:
        emphasis.add_run("Liên hệ trực tiếp nội dung slide với pipeline thực tế và nêu rõ vai trò của bước này trong việc chọn sản phẩm cuối cùng.")

    if index < len(slides):
        doc.add_paragraph("—" * 18).alignment = WD_ALIGN_PARAGRAPH.CENTER

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Aura Fashion • AI Recommendation • Giải thích từng slide")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(110, 120, 135)

doc.save(OUTPUT)
print(OUTPUT)
