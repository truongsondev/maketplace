from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "Tom_Tat_AI_Recommendation_Aura.docx"


slides = [
    ("AI Recommendation", "Giới thiệu hệ thống gợi ý Aura Fashion và mục tiêu giải thích bằng kỹ thuật vì sao một sản phẩm cụ thể được xếp hạng số 1."),
    ("Vì sao sản phẩm này được gợi ý?", "Kết quả gợi ý được hình thành từ bằng chứng hành vi người dùng, đặc trưng nội dung sản phẩm và các bước hợp nhất điểm, lọc, xếp hạng."),
    ("Kiến trúc triển khai", "Storefront phát sự kiện; RabbitMQ tiếp nhận; Node.js điều phối candidate, luật và cache; FastAPI xử lý embedding/hybrid; pgvector tìm kiếm cosine; Redis và MySQL lưu tín hiệu, artifact."),
    ("Dữ liệu đầu vào", "Mỗi RecommendationEvent chứa loại hành vi, user/session, sản phẩm, truy vấn tìm kiếm, placement, thời gian và dedupeKey để cá nhân hóa, đo lường và chống ghi trùng."),
    ("Ingestion qua RabbitMQ", "RabbitMQ tách độ trễ giao diện khỏi xử lý nền. Consumer persist sự kiện, cập nhật feature Redis, xóa cache; message retry tối đa ba lần rồi vào DLQ và được bảo vệ bằng idempotency."),
    ("Trọng số hành vi", "Purchase có giá trị cao hơn add-to-cart, favorite và view. Mỗi feed dùng trọng số hơi khác theo mục tiêu; đây là heuristic kinh doanh, chưa phải trọng số được mô hình học."),
    ("Offline refresh", "Khoảng 15 phút một lần, hệ thống đọc catalog và sự kiện, tính co-occurrence từ đơn hàng, huấn luyện preference/popularity và tạo embedding sản phẩm để phục vụ nhanh."),
    ("Document sản phẩm", "Tên, mô tả, danh mục, tag và thuộc tính được ghép thành văn bản, tokenize, encode thành vector 384 chiều, chuẩn hóa L2 rồi lưu. Chất lượng đầu vào quyết định chất lượng semantic retrieval."),
    ("Encoder và fallback", "Ưu tiên all-MiniLM-L6-v2 để hiểu ngữ nghĩa; khi không tải được model, dùng hashing 384 chiều. Fallback giữ dịch vụ hoạt động nhưng không hiểu tốt từ đồng nghĩa và ngữ cảnh."),
    ("Cosine similarity", "pgvector dùng khoảng cách cosine để tìm sản phẩm gần vector ngữ cảnh. Với vector đã chuẩn hóa, cosine gần tương đương tích vô hướng; IVFFlat hỗ trợ tìm kiếm nhanh hơn full scan."),
    ("Context vector", "Gu hiện tại của người dùng được biểu diễn bằng trung bình vector của tối đa các sản phẩm tương tác gần đây, sau đó chuẩn hóa. Khi không có context, hệ thống phải dựa nhiều hơn vào popularity/fallback."),
    ("Công thức hybrid", "Điểm AI = 0,75 × cosine + 0,25 × log(1 + popularity). Semantic là thành phần chính theo thiết kế, nhưng popularity chưa chuẩn hóa có thể vượt thang 0–1 và làm đảo hạng."),
    ("Candidate generation", "Ứng viên đến từ nhiều nguồn: pgvector, co-occurrence, hot score, trending, search intent, category, hành vi gần đây và catalog fallback. Hệ thống ưu tiên recall rồi mới lọc và lấy top-K."),
    ("Feed Home", "Home kết hợp hot Redis, trending 30 ngày, ý định tìm kiếm và hành vi session. Nó vẫn tạo được kết quả cho khách ẩn danh; nếu thiếu context thì ưu tiên sản phẩm sale, có ảnh và mới."),
    ("Feed Product Detail", "Sản phẩm tương tự được lấy từ vector semantic, đồng xuất hiện trong đơn và cùng category/type. Sản phẩm đang xem luôn bị loại khỏi kết quả."),
    ("Feed Cart", "Giỏ hàng tạo vector ngữ cảnh trung bình từ nhiều sản phẩm. AI, co-occurrence và category sinh candidate cross-sell; mọi sản phẩm đang trong giỏ đều bị loại."),
    ("Feed Personalized", "Đây là pipeline đầy đủ nhất: hành vi 45 ngày, search intent, AI, co-occurrence, category và fallback; đồng thời loại context, giỏ hàng và sản phẩm đã mua trong 90 ngày."),
    ("Search intent", "Các truy vấn gần đây được xếp theo tần suất và độ mới rồi tìm bằng SQL LIKE trên tên/mô tả. Cách này dễ giải thích nhưng chưa hiểu typo, từ đồng nghĩa hay ngữ nghĩa câu."),
    ("Co-occurrence", "Self-join order_items đếm số đơn có đồng thời hai sản phẩm, tạo tín hiệu item-item minh bạch. Nhược điểm là popularity bias vì chưa chuẩn hóa theo độ phổ biến từng sản phẩm."),
    ("Merge score trong Node", "Khi một sản phẩm đến từ nhiều nguồn, giữ điểm lớn nhất và cộng 20% từng điểm phụ. Vì các nguồn dùng thang điểm khác nhau, category hoặc count có thể lấn át cosine."),
    ("Bộ lọc và guardrail", "Hệ thống loại sản phẩm trùng context, đang trong giỏ, đã mua gần đây, bị xóa hoặc thiếu dữ liệu card. Cần bổ sung lọc trạng thái ACTIVE và tồn kho khả dụng."),
    ("Ví dụ tính AI score", "Trong ví dụ, sản phẩm có cosine thấp hơn vẫn đứng đầu vì popularity cao. Điều này chứng minh popularity cần được normalize nếu muốn tỷ lệ 75/25 phản ánh đúng ý nghĩa."),
    ("Ví dụ sau merge", "AI winner chưa chắc là kết quả cuối. Điểm category và co-occurrence ở tầng Node có thể đảo thứ hạng sau hợp nhất, trước khi candidate hợp lệ được hiển thị."),
    ("Luồng end-to-end", "Một sản phẩm đi qua sáu bước: tạo document, embedding, tạo context, cosine retrieval, hybrid scoring, merge/filter. Nó đứng số 1 vì sống sót và đạt điểm cao qua toàn pipeline."),
    ("Cache hai tầng", "Redis là fast path; MySQL RecommendationCache là fallback bền vững. Cache key phụ thuộc feed và context, với TTL khoảng 10–30 phút tùy placement."),
    ("Cache invalidation", "Sự kiện mới chủ động xóa cache liên quan đến session/user để feed đổi ngay, không chờ hết TTL. Redis KEYS nên được thay bằng SCAN hoặc versioned key khi quy mô lớn."),
    ("Fault tolerance", "Nếu AI timeout hoặc lỗi, Node bắt exception và tiếp tục bằng co-occurrence, category hoặc catalog; RabbitMQ có retry/DLQ và DB cache hỗ trợ phục hồi. Cần timeout rõ cho request AI."),
    ("Observability", "Prometheus theo dõi số event, cache hit, thời gian tạo feed, độ trễ AI và số lần fallback. Nên bổ sung số candidate theo nguồn, zero-result, phân phối score và CTR/CVR theo model."),
    ("Đánh giá chất lượng", "Offline dùng Recall@K, NDCG@K, MRR và coverage; online đo CTR, add-to-cart, purchase và doanh thu. Guardrail gồm latency, hết hàng, diversity, novelty và độ tập trung popularity."),
    ("A/B testing", "Gán user/session ổn định vào control/treatment, ghi impression cùng variant, theo dõi outcome và so uplift có ý nghĩa thống kê. Schema experiment có sẵn nhưng serving chưa nối hoàn chỉnh."),
    ("Technical debt", "Các vấn đề chính gồm score khác thang, body profile/user embedding chưa dùng, thiếu lọc stock/status, hashing không ổn định, IVFFlat chưa tune, exclusion xung đột và search chưa embedding."),
    ("Roadmap kỹ thuật", "Ưu tiên guardrail tồn kho, chuẩn hóa điểm, rank fusion, diversity reranking, query embedding và A/B testing; sau khi log đủ tốt mới tiến tới Learning-to-Rank bằng LightGBM/XGBoost."),
    ("Kịch bản demo", "Demo nên theo dõi event và hot score, refresh artifact, truy vấn cosine, gọi feed xem source/reason, kiểm chứng cache hit rồi phát event mới để thấy cache bị xóa và ranking thay đổi."),
    ("Kết luận", "Recommendation là chuỗi quyết định có thể kiểm chứng: retrieval tăng recall, hybrid và merge tạo thứ hạng, guardrail bảo vệ kết quả, còn event–cache–refresh tạo vòng lặp học liên tục."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.color.rgb = RGBColor(29, 53, 87)
styles["Heading 1"].font.color.rgb = RGBColor(29, 53, 87)
styles["Heading 2"].font.color.rgb = RGBColor(67, 97, 238)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("AURA FASHION").bold = True
p.runs[0].font.size = Pt(14)
p.runs[0].font.color.rgb = RGBColor(67, 97, 238)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("TÓM TẮT BÀI THUYẾT TRÌNH\nAI RECOMMENDATION")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Từ tín hiệu hành vi đến một sản phẩm được xếp hạng #1")
run.italic = True
run.font.size = Pt(13)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Tài liệu tóm tắt 34 slide • Aura Fashion Marketplace • 2026").font.color.rgb = RGBColor(90, 100, 115)

doc.add_page_break()
doc.add_heading("1. Tóm tắt tổng quan", level=1)
doc.add_paragraph(
    "Hệ thống AI Recommendation của Aura Fashion không phụ thuộc vào một mô hình duy nhất. "
    "Nó là pipeline lai kết hợp tín hiệu hành vi, nội dung sản phẩm, embedding 384 chiều, "
    "tìm kiếm cosine bằng pgvector, popularity, co-occurrence, luật nghiệp vụ, cache và fallback."
)
for item in [
    "Dữ liệu hành vi được thu nhận bất đồng bộ qua RabbitMQ, lưu vào MySQL/Redis và dùng để cập nhật feature cũng như xóa cache.",
    "Sản phẩm được chuyển thành vector semantic; vector ngữ cảnh người dùng là trung bình các vector sản phẩm tương tác gần đây.",
    "AI service tính điểm hybrid 0,75 × cosine + 0,25 × log(1 + popularity).",
    "Node.js hợp nhất candidate từ nhiều nguồn theo quy tắc max + 20% điểm phụ, sau đó loại sản phẩm không phù hợp và lấy top-K.",
    "Điểm cần cải thiện quan trọng nhất là chuẩn hóa score giữa các nguồn, bổ sung guardrail tồn kho/trạng thái và đo chất lượng bằng A/B test.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("2. Pipeline cốt lõi", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Shading Accent 1"
headers = ["Giai đoạn", "Xử lý chính", "Kết quả"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    set_cell_shading(table.rows[0].cells[i], "4357EE")
    for r in table.rows[0].cells[i].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.bold = True
pipeline = [
    ("1. Thu nhận", "UI phát VIEW/CART/FAVORITE/PURCHASE qua RabbitMQ", "Event tin cậy, chống trùng"),
    ("2. Biểu diễn", "Ghép metadata và encode MiniLM 384D", "Vector sản phẩm"),
    ("3. Tạo ngữ cảnh", "Lấy trung bình vector tương tác gần đây", "Vector gu hiện tại"),
    ("4. Sinh candidate", "pgvector, co-occurrence, hot, search, category", "Tập ứng viên recall cao"),
    ("5. Xếp hạng", "Hybrid AI rồi merge đa nguồn", "Điểm cuối trong Node"),
    ("6. Guardrail", "Exclude, unique, soft-delete, card lookup", "Top-K hợp lệ"),
    ("7. Serving", "Redis/MySQL cache và fallback", "Feed nhanh, chịu lỗi"),
]
for row in pipeline:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

doc.add_page_break()
doc.add_heading("3. Tóm tắt theo từng slide", level=1)
for index, (name, summary) in enumerate(slides, 1):
    doc.add_heading(f"Slide {index:02d} — {name}", level=2)
    doc.add_paragraph(summary)

doc.add_heading("4. Thông điệp kết luận", level=1)
quote = doc.add_paragraph()
quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = quote.add_run(
    "Một sản phẩm đứng #1 không phải vì API tình cờ trả về nó, mà vì sản phẩm đó có bằng chứng phù hợp nhất "
    "sau toàn bộ quá trình retrieval, scoring, merge, exclusion và guardrail."
)
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(29, 53, 87)

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Aura Fashion • AI Recommendation Summary").font.size = Pt(8)

doc.save(OUTPUT)
print(OUTPUT)
